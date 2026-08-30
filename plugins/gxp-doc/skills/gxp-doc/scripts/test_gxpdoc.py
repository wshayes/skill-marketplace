"""Self-check for gxpdoc.py. Run:

    uv run --with python-docx python .claude/skills/gxp-doc/scripts/test_gxpdoc.py
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches  # noqa: E402

import gxpdoc  # noqa: E402
from gxpdoc import GxpDoc  # noqa: E402


def build():
    doc = GxpDoc(doc_id="SOP-999", title="Self Check Procedure",
                 doc_type="sop", version="1.0", company="Acme Pharma")
    doc.toc()
    doc.h1("Purpose")
    doc.clause(2, "To exercise every construct the builder can emit.")
    doc.h1("Procedure")
    doc.clause(2, "The Document Owner drafts the document.")
    doc.clause(3, "A wrapped clause " + "with enough text to wrap " * 8)
    doc.clause(4, "Change management shall include the following:")
    doc.bullet("Description of the change and the justification")
    doc.table(["Prefix", "Type", "Example"],
              [("SOP", "Standard Operating Procedure", "SOP-003")])
    doc.references({
        "QMS Documents": [("SOP-003", "Change Control Procedure")],
        "External Standards": [("21 CFR Part 11", "Electronic Records")],
    })
    doc.revision_history([
        ("1.0", "2026-05-22", "Quality Team", "Initial release", ""),
        ("2.0", "", "Quality Team", "Restyled to house standard", "CC-2026-014"),
    ])
    doc.approval([("Reviewer", "William Hayes", "Head of IT")])
    return doc


def runs(el):
    return el.findall(qn("w:r"))


def main():
    built = build()
    with tempfile.TemporaryDirectory() as tmp:
        path = built.save(Path(tmp) / "SOP-999_Self_Check_Procedure_v1.0.docx")
        doc = Document(str(path))

        section = doc.sections[0]
        assert section.page_width == Inches(8.5), "page must be US Letter"
        assert section.page_height == Inches(11)
        assert section.left_margin == Inches(1)

        # -- header block -------------------------------------------------
        header = section.header
        assert header.paragraphs[0].text == "Acme Pharma", \
            header.paragraphs[0].text
        assert len(header.tables) == 1, "header carries exactly one table"
        table = header.tables[0]
        assert table.rows[0].cells[0].text == "STANDARD OPERATING PROCEDURE"
        assert [c.text for c in table.rows[1].cells] == [
            "Document ID", "Document Description", "Version No."]
        assert [c.text for c in table.rows[2].cells] == [
            "SOP-999", "SELF CHECK PROCEDURE", "1.0"]

        borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        assert borders is not None, "header table needs explicit borders"
        assert borders.find(qn("w:top")).get(qn("w:sz")) == "12", "outer 1.5pt"
        assert borders.find(qn("w:insideH")).get(qn("w:sz")) == "6", "inner 0.75pt"
        assert borders.find(qn("w:top")).get(qn("w:color")) == "000000", "black"
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.tcPr
                shd = None if tc_pr is None else tc_pr.find(qn("w:shd"))
                assert shd is None, "header table carries no shading"

        # -- footer: live page fields, centered ---------------------------
        footer_xml = section.footer.paragraphs[0]._p.xml
        assert "PAGE" in footer_xml and "NUMPAGES" in footer_xml
        assert 'w:fldCharType="begin"' in footer_xml, \
            "page numbers must be live fields, not frozen text"
        stops = [t.position for t in
                 section.footer.paragraphs[0].paragraph_format.tab_stops]
        assert Inches(3.25) in stops, "Page X of Y is centered"

        # -- nothing but the contents field precedes section 1 -------------
        assert doc.tables[0].rows[0].cells[0].text != "Field", \
            "no metadata block: the page header carries identity"

        # -- table of contents --------------------------------------------
        body_xml = doc.element.body.xml
        assert 'TOC \\o "1-1"' in body_xml, "level-1 table of contents"

        # -- clause ladder --------------------------------------------------
        levels = {}
        for para in doc.paragraphs:
            if para.style.name.startswith("GxP L"):
                levels.setdefault(para.style.name, []).append(para.text)

        assert levels["GxP L1"][0].startswith("1.\tPURPOSE"), levels["GxP L1"][0]
        assert levels["GxP L1"] == [
            "1.\tPURPOSE", "2.\tPROCEDURE", "3.\tREFERENCES",
            "4.\tREVISION HISTORY", "5.\tAPPROVAL",
        ], levels["GxP L1"]
        assert levels["GxP L2"][1].startswith("2.1.\t"), levels["GxP L2"]
        assert levels["GxP L3"][0].startswith("2.1.1.\t"), levels["GxP L3"]
        assert levels["GxP L4"][0].startswith("2.1.1.1.\t"), levels["GxP L4"]

        for level, (number_x, text_x) in gxpdoc.INDENT.items():
            fmt = doc.styles[f"GxP L{level}"].paragraph_format
            assert fmt.left_indent == gxpdoc._inches(text_x), level
            assert fmt.first_line_indent == gxpdoc._inches(number_x - text_x), level

        assert doc.styles["GxP L1"].font.size.pt == 13
        assert doc.styles["GxP L1"].font.bold is True
        assert doc.styles["GxP L1"].font.all_caps is True
        for level in (2, 3, 4):
            assert doc.styles[f"GxP L{level}"].font.size.pt == 11, level
            assert not doc.styles[f"GxP L{level}"].font.bold, level

        # numbering is literal text, never a Word auto-list
        for para in doc.paragraphs:
            if para.style.name.startswith("GxP L"):
                p_pr = para._p.pPr
                assert p_pr is None or p_pr.find(qn("w:numPr")) is None, \
                    "clause numbers must be literal text, not auto-numbering"

        # -- body tables ----------------------------------------------------
        for table in doc.tables:
            head = table.rows[0]
            shd = head.cells[0]._tc.tcPr.find(qn("w:shd"))
            assert shd is not None and shd.get(qn("w:fill")) == "F2F2F2", \
                "body table header rows are shaded F2F2F2"
            assert head._tr.trPr.find(qn("w:tblHeader")) is not None, \
                "body table header rows repeat across page breaks"
            borders = table._tbl.tblPr.find(qn("w:tblBorders"))
            assert borders.find(qn("w:top")).get(qn("w:color")) == "D9D9D9"

        headers = [[c.text for c in t.rows[0].cells] for t in doc.tables]
        assert ["Role", "Name", "Title", "Signature", "Date"] in headers
        assert ["Version", "Date", "Author", "Description of Change",
                "Change Control Ref"] in headers
        assert ["Document ID", "Title"] in headers

        # revision history newest first
        rev = next(t for t in doc.tables if t.rows[0].cells[0].text == "Version")
        assert [r.cells[0].text for r in rev.rows[1:]] == ["2.0", "1.0"]
        # approval dates blank until signed
        appr = next(t for t in doc.tables if t.rows[0].cells[0].text == "Role")
        assert appr.rows[1].cells[4].text == ""

        # -- one font, all black ---------------------------------------------
        for para in doc.paragraphs:
            for run in para.runs:
                assert run.font.name == gxpdoc.FONT, (para.text, run.font.name)
                colour = run.font.color
                assert colour.rgb is None or str(colour.rgb) == "000000", para.text

    # -- constructor guards ---------------------------------------------
    for bad_version in ("DRAFT", "draft", "1", "v1.0"):
        try:
            GxpDoc("SOP-999", "X", "sop", bad_version)
        except ValueError:
            pass
        else:  # pragma: no cover
            raise AssertionError(f"version {bad_version!r} should be rejected")

    assert not hasattr(GxpDoc, "control_block"), \
        "the metadata block is gone; the page header carries identity now"

    # The skill is company-agnostic: the wordmark defaults to a placeholder
    # that the linter refuses to let through into a signature copy.
    assert gxpdoc.COMPANY == gxpdoc.COMPANY_PLACEHOLDER == "<Company Name>"
    default = GxpDoc("SOP-999", "X", "sop", "1.0")
    assert default.doc.sections[0].header.paragraphs[0].text == "<Company Name>"

    with tempfile.TemporaryDirectory() as tmp:
        for bad_name in ("SOP-999.docx", "SOP-999_X_v2.0.docx",
                         "SOP-998_X_v1.0.docx"):
            try:
                build().save(Path(tmp) / bad_name)
            except ValueError:
                pass
            else:  # pragma: no cover
                raise AssertionError(f"filename {bad_name!r} should be rejected")
        # DRAFT suffix is the only place draft state may appear
        build().save(Path(tmp) / "SOP-999_Self_Check_Procedure_v1.0_DRAFT3.docx")

    # -- fields Word must compute on open --------------------------------
    # A field carries a cached result and python-docx has no layout engine to
    # compute one, so without w:updateFields every document opens showing the
    # placeholder text instead of contents and page numbers.
    settings = built.doc.settings.element.findall(qn("w:updateFields"))
    assert len(settings) == 1, "exactly one w:updateFields flag"
    assert settings[0].get(qn("w:val")) == "true", \
        "Word must be told to refresh fields on open, or the TOC shows its " \
        "placeholder and the footer reads 'Page 1 of 1'"

    # -- clause styles carry an outline level ------------------------------
    # TOC \o "1-1" collects by outline level, not by style name. A style based
    # on Normal inherits "body text", which makes the headings invisible to
    # the field.
    for level in (1, 2, 3, 4):
        ppr = built.doc.styles[f"GxP L{level}"].element.find(qn("w:pPr"))
        node = None if ppr is None else ppr.find(qn("w:outlineLvl"))
        assert node is not None, f"GxP L{level} has no outline level"
        assert node.get(qn("w:val")) == str(level - 1), \
            f"GxP L{level} must sit at outline level {level - 1}"

    # -- the linter agrees the builder's own output is house style ---------
    # Guards both directions of G/9: SOP-003 is cited only from a table cell
    # here, which justifies its References row without demanding new ones.
    import check_doc  # noqa: PLC0415 -- optional, keeps gxpdoc importable alone

    with tempfile.TemporaryDirectory() as tmp:
        path = build().save(Path(tmp) / "SOP-999_Self_Check_Procedure_v1.0.docx")
        findings = check_doc.check(path)
        # The sample above is a deliberately partial SOP, so D/5 fires on the
        # sections it omits. Everything the References rules touch must be
        # silent: SOP-003 is cited only from a table cell here, and that has
        # to satisfy its References row without demanding a new one.
        noisy = [i for i in findings.items
                 if i["rule"].startswith(("F/", "G/"))]
        assert not noisy, f"cross-reference rules must be silent, got {noisy}"

    # A two-digit sequence suffix is a valid identifier (style guide A):
    # WI-CSA-01 and FRM-CSA-07 are cited that way throughout a validation set.
    for identifier in ("WI-CSA-01", "FRM-CSA-07", "SOP-003", "SUP-2026-001",
                       "GXP-URS-001", "TQ-001"):
        text = f"{identifier} (Some Name)"
        match = check_doc.ID_RE.search(text)
        assert match and match.group(0) == identifier, \
            f"{identifier} must parse as one identifier, got {match}"
        assert check_doc.RESOLVED_RE.match(text[match.end():]), \
            f"{identifier} followed by a name must count as resolved"

    print("gxpdoc self-check: all assertions passed")


if __name__ == "__main__":
    main()
