"""Lint a controlled document against the GxP style guide.

    uv run --with python-docx --with openpyxl python \
        .claude/skills/gxp-doc/scripts/check_doc.py docs_qms/*.docx

Exits non-zero if any error-level finding is reported. --json emits machine
readable output. Rule ids map to the components in ../references/style-guide.md.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.path.insert(0, str(Path(__file__).parent))
from gxpdoc import (  # noqa: E402
    BANNER,
    COMPANY_PLACEHOLDER,
    FILENAME_RE,
    FONT,
    FORBIDDEN_CONTROL_ROWS,
    LEGACY_CONTROL_ROWS,
    VERSION_RE,
)

ERROR, WARN = "error", "warning"

ID_RE = re.compile(
    r"\b(?:(?:QM|POL|SOP|WI|FRM|LOG|SUP|TQ|CC)-[0-9A-Z]{2,4}(?:-[0-9]{2,3})?"
    r"|GXP-[A-Z]{2,4}(?:-[A-Z]{2,5})?-[0-9]{3})\b"
)
#: An identifier is resolved if a name follows it, optionally after a section ref.
RESOLVED_RE = re.compile(r"^(?:\s*§[\d.]+)?\s*(?:\(|—\s|-\s)")
#: Numbering-convention placeholders ("SUP-YYYY-NNN"), not real citations.
PLACEHOLDER_ID_RE = re.compile(r"^(?:[A-Z]{2,4}-)?(?:YYYY|NNN|XXX|N{2,})\b")
LABEL_RE = re.compile(r"^\s*([A-Z][\w &/,'-]{1,60}):\s+\S")
CLAUSE_RE = re.compile(r"^(\d+(?:\.\d+)*)\.\t?\s*(.*)$", re.DOTALL)
NON_ISO_DATE_RE = re.compile(
    r"\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}"
    r"|\d{1,2}\s?(?:JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)[A-Z]*\s?\d{4})\b",
    re.IGNORECASE,
)
PLACEHOLDER_RE = re.compile(r"<[a-z][^>]{2,40}>|\bTBD\b|\bTODO\b", re.IGNORECASE)

VAGUE_TERMS = [
    "appropriate", "adequate", "periodically", "as needed", "if necessary",
    "as required", "timely", "regularly",
]
UK_SPELLINGS = [
    "organisation", "organisations", "prioritise", "authorise", "authorised",
    "recognise", "centre", "analyse", "catalogue", "behaviour", "colour",
    "favour", "labour", "minimise", "utilise", "fulfil",
]

SOP_SKELETON = [
    "PURPOSE", "SCOPE", "DEFINITIONS", "RESPONSIBILITIES", "PROCEDURE",
    "RECORDS AND RETENTION", "REFERENCES", "REVISION HISTORY", "APPROVAL",
]
TAIL_SKELETON = ["REFERENCES", "REVISION HISTORY", "APPROVAL"]

PREFIX_TO_TYPE = {
    "QM": "qm", "POL": "policy", "SOP": "sop", "WI": "wi",
    "FRM": "form", "LOG": "log", "GXP": "protocol", "TQ": "protocol",
    "SUP": "form",
}


class Findings:
    def __init__(self, path):
        self.path = Path(path)
        self.items = []

    def add(self, level, rule, message, where=""):
        self.items.append(
            {"level": level, "rule": rule, "message": message, "where": where}
        )

    error = lambda self, *a, **k: self.add(ERROR, *a, **k)  # noqa: E731
    warn = lambda self, *a, **k: self.add(WARN, *a, **k)  # noqa: E731

    @property
    def errors(self):
        return [i for i in self.items if i["level"] == ERROR]


def blocks(container, doc):
    """Yield ('p', Paragraph) / ('t', Table) in true document order."""
    for child in container.iterchildren():
        if child.tag == qn("w:p"):
            yield "p", Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield "t", Table(child, doc)


def rows_of(table):
    return [[c.text.strip() for c in r.cells] for r in table.rows]


def style_name(paragraph):
    try:
        return paragraph.style.name or ""
    except AttributeError:
        return ""


def is_level1(paragraph):
    """A top-level section heading.

    Legacy documents number some section headings and not others; the
    numbering itself is checked separately by the clause-ladder rule, so
    recognising a section must not depend on it.
    """
    return style_name(paragraph) in {"GxP L1", "Heading 1", "Heading 2"}


def heading_text(paragraph):
    text = paragraph.text.strip()
    return re.sub(r"^\d+(?:\.\d+)*\.?\s*", "", text).strip().upper()


# ---------------------------------------------------------------------------
# Rules
# ---------------------------------------------------------------------------


def check_header(doc, f, doc_id, version, is_draft=True):
    """B — wordmark + 3-column banner table with black rules, no shading."""
    header = doc.sections[0].header
    if not header.tables:
        f.error("B/1", "no header table: the standardized 3-column banner "
                       "block is missing")
        return
    table = header.tables[0]
    grid = rows_of(table)
    if len(grid) != 3 or len(grid[0]) != 3:
        f.error("B/1", f"header table is {len(grid)}x{len(grid[0])}, expected 3x3")
        return

    banner = grid[0][0].strip().upper()
    if banner not in set(BANNER.values()):
        f.error("B/1", f"header banner {banner!r} is not one of "
                       f"{sorted(set(BANNER.values()))}")
    labels = [c.strip() for c in grid[1]]
    if labels != ["Document ID", "Document Description", "Version No."]:
        f.error("B/1", f"header column labels are {labels}")
    if doc_id and grid[2][0].strip() != doc_id:
        f.error("B/1", f"header says {grid[2][0]!r}, filename says {doc_id!r}")
    if version and grid[2][2].strip() != version:
        f.error("B/1", f"header version {grid[2][2]!r} != filename {version!r}")

    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    if borders is None:
        f.error("B/1", "header table has no explicit borders (needs solid black "
                       "1.5pt outer / 0.75pt inner)")
    else:
        top, inside = borders.find(qn("w:top")), borders.find(qn("w:insideH"))
        if top is None or top.get(qn("w:sz")) != "12":
            f.error("B/1", "header table outer border is not 1.5pt")
        if inside is None or inside.get(qn("w:sz")) != "6":
            f.error("B/1", "header table inner border is not 0.75pt")
        if top is not None and top.get(qn("w:color")) not in (None, "000000"):
            f.error("B/1", "header table borders must be black")
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.tcPr
            if tc_pr is not None and tc_pr.find(qn("w:shd")) is not None:
                f.error("B/1", "header table must not be shaded")
                break

    wordmark = header.paragraphs[0].text.strip() if header.paragraphs else ""
    if not wordmark and not (header.paragraphs and header.paragraphs[0].runs):
        f.warn("B/1", "no company wordmark above the header table")
    elif wordmark == COMPANY_PLACEHOLDER:
        level = f.warn if is_draft else f.error
        level("B/1", f"company wordmark is still {COMPANY_PLACEHOLDER}")


def check_footer(doc, f):
    """B — centered Page X of Y built from live fields."""
    footer = doc.sections[0].footer
    xml = "".join(p._p.xml for p in footer.paragraphs)
    if "PAGE" not in xml or "NUMPAGES" not in xml:
        f.error("B/2", "footer has no Page X of Y")
        return
    if 'w:fldCharType="begin"' not in xml:
        f.error("B/2", "footer page numbers are frozen text, not live fields — "
                       "this is what makes documents print 'Page  of '")
    centered = any(
        abs(t.position.inches - 3.25) < 0.05
        for p in footer.paragraphs for t in p.paragraph_format.tab_stops
    )
    if not centered:
        f.warn("B/2", "Page X of Y is not on the centered tab stop")


def check_page_setup(doc, f):
    """B — US Letter, 1in margins."""
    s = doc.sections[0]
    if (round(s.page_width.inches, 2), round(s.page_height.inches, 2)) != (8.5, 11.0):
        f.error("B/3", f"page is {s.page_width.inches:.2f}x"
                       f"{s.page_height.inches:.2f}in, expected US Letter 8.5x11")
    for side in ("left", "right", "top", "bottom"):
        if abs(getattr(s, f"{side}_margin").inches - 1.0) > 0.02:
            f.error("B/3", f"{side} margin is "
                           f"{getattr(s, f'{side}_margin').inches:.2f}in, expected 1in")


KNOWN_FIELDS = {k.lower() for k in LEGACY_CONTROL_ROWS} | FORBIDDEN_CONTROL_ROWS | {
    "document id", "environment", "references", "form used", "parent procedure",
}


def find_control_block(ordered):
    """The Field|Details table, or a legacy headerless key/value table.

    Several validation documents open with a bare two-column table whose
    first row is already data ("Document ID | GXP-OQ-001"). Skipping those
    would hide exactly the Status/Date rows this linter exists to catch.
    """
    fallback = None
    for kind, block in ordered:
        if kind != "t":
            continue
        grid = rows_of(block)
        if not grid:
            continue
        if [c.lower() for c in grid[0][:2]] in (["field", "details"],
                                                ["field", "value"]):
            return block
        if fallback is None and len(grid[0]) == 2:
            keys = {r[0].strip().lower() for r in grid if r[0].strip()}
            if len(keys & KNOWN_FIELDS) >= 2:
                fallback = block
    return fallback


def check_no_control_block(table, f):
    """C — a metadata block is redundant with the repeating page header."""
    if table is None:
        return {}
    grid = rows_of(table)
    body = grid[1:] if [c.lower() for c in grid[0][:2]] in (
        ["field", "details"], ["field", "value"]) else grid
    fields = {r[0].strip(): r[1].strip() for r in body if r[0].strip()}
    f.error("C/4", "metadata block before the first section: the page header "
                   "already carries ID, description and version, and nothing "
                   "else belongs in the document — delete it")
    for key, value in fields.items():
        if key.lower() in FORBIDDEN_CONTROL_ROWS:
            f.error("C/4a", f"{key} = {value!r} must move to the Document "
                            "Master List; a document holds no status and no date")
    return fields


def header_identity(doc):
    """(document id, version) as the repeating page header states them."""
    header = doc.sections[0].header
    if not header.tables:
        return "", ""
    grid = rows_of(header.tables[0])
    if len(grid) < 3 or len(grid[2]) < 3:
        return "", ""
    return grid[2][0].strip(), grid[2][2].strip()


def check_version_and_filename(path, version, doc_id, f):
    """A — version is a real number; filename carries the draft state."""
    if version.upper().startswith("DRAFT") or version.upper() == "DRAFT":
        f.error("A/4b", "version reads DRAFT — the version field always carries "
                        "the next real number; draft state belongs in the filename")
    elif version and not VERSION_RE.match(version):
        f.error("A/4b", f"version {version!r} is not X.Y")

    match = FILENAME_RE.match(path.name)
    if not match:
        f.error("A/4c", "filename does not match "
                        "<ID>_<Title_Snake>_v<X.Y>[_DRAFT<N>].docx")
        return False
    if version and match.group("version") != version:
        f.error("A/4c", f"filename says v{match.group('version')} but the "
                        f"header says {version}")
    if doc_id and match.group("id") != doc_id:
        f.error("A/4c", f"filename says {match.group('id')} but the header "
                        f"says {doc_id}")
    return "_DRAFT" in path.name


def check_skeleton(sections, doc_type, f):
    """D — canonical section order per document type."""
    expected = SOP_SKELETON if doc_type in ("sop", "policy") else TAIL_SKELETON
    missing = [s for s in expected if s not in sections]
    if missing:
        f.error("D/5", f"missing sections: {', '.join(missing)}")
    if doc_type in ("protocol", "report"):
        # Validation documents deliberately front-load Revision History and
        # Approval -- sign before execute. Presence is required, order is not.
        return
    present = [s for s in sections if s in expected]
    ordered = [s for s in expected if s in sections]
    if present != ordered:
        f.error("D/5", f"section order is {present}, expected {ordered}")


def check_tables(ordered, f):
    """E — gray shaded, repeating, lightly ruled header rows."""
    for kind, table in ordered:
        if kind != "t":
            continue
        head = table.rows[0]
        label = head.cells[0].text.strip()[:24] or "(blank)"
        tc_pr = head.cells[0]._tc.tcPr
        shd = None if tc_pr is None else tc_pr.find(qn("w:shd"))
        if shd is None or shd.get(qn("w:fill")) not in ("F2F2F2",):
            f.error("E/6", "table header row is not shaded F2F2F2", label)
        if head._tr.trPr is None or head._tr.trPr.find(qn("w:tblHeader")) is None:
            f.error("E/6", "table header row does not repeat across page breaks",
                    label)
        borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        if borders is None:
            f.error("E/6", "table has no explicit borders", label)


def check_table_candidates(ordered, f):
    """E — table-shaped data must be a table, not a bullet list."""
    run, first = [], None
    for kind, block in ordered:
        if kind == "p" and LABEL_RE.match(block.text.strip()):
            if not run:
                first = block.text.strip()[:50]
            run.append(block.text)
            continue
        if len(run) >= 3:
            f.error("E/6a", f"{len(run)} consecutive 'Label: value' paragraphs "
                            "should be a table", first)
        run, first = [], None
    if len(run) >= 3:
        f.error("E/6a", f"{len(run)} consecutive 'Label: value' paragraphs "
                        "should be a table", first)


def check_typography(doc, f):
    """J — Times New Roman everywhere, all text black."""
    fonts, colours = set(), set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                fonts.add(run.font.name)
            if run.font.color is not None and run.font.color.rgb is not None:
                colours.add(str(run.font.color.rgb))
    wanted = {"Normal", "Heading 1", "Heading 2", "Heading 3"}
    for style in doc.styles:
        if style.name in wanted and style.font.name:
            fonts.add(style.font.name)
    stray_fonts = sorted(x for x in fonts if x != FONT)
    if stray_fonts:
        f.error("J/6b", f"fonts other than {FONT}: {', '.join(stray_fonts)}")
    stray_colours = sorted(c for c in colours if c != "000000")
    if stray_colours:
        f.error("J/6b", f"non-black text: {', '.join(stray_colours)}")


def check_clause_ladder(doc, ordered, f):
    """J — indented, literally numbered clause levels with trailing periods."""
    styled = [b for k, b in ordered if k == "p" and style_name(b).startswith("GxP L")]
    if not styled:
        f.warn("J/6c", "no GxP L1–L4 clause styles: document predates the "
                       "house heading scheme")
    counters = {}
    for para in styled:
        level = int(style_name(para)[-1])
        text = para.text.strip()
        match = CLAUSE_RE.match(text)
        if not match:
            f.error("J/6c", f"level {level} clause is not numbered", text[:50])
            continue
        number = match.group(1)
        if text[len(number)] != ".":
            f.error("J/6c", f"clause {number} has no trailing period", text[:50])
        if number.count(".") + 1 != level:
            f.error("J/6c", f"clause {number} is styled as level {level}",
                    text[:50])
        parent = number.rsplit(".", 1)[0] if "." in number else ""
        want = counters.get(parent, 0) + 1
        got = int(number.rsplit(".", 1)[-1])
        if got != want:
            f.error("J/6c", f"clause numbering jumps: expected "
                            f"{parent + '.' if parent else ''}{want}, got {number}")
        counters[parent] = got
        counters = {k: v for k, v in counters.items()
                    if not k.startswith(number + ".")}
        p_pr = para._p.pPr
        if p_pr is not None and p_pr.find(qn("w:numPr")) is not None:
            f.error("J/6c", f"clause {number} uses Word auto-numbering; numbers "
                            "must be literal text so §-citations stay stable")
        if level == 1 and text != text.upper():
            f.error("J/6c", "level 1 headings are ALL CAPS", text[:50])


def check_toc_and_title(ordered, doc, f):
    """J — a contents field, then section 1. Nothing else may precede it."""
    body_xml = doc.element.body.xml
    if 'TOC \\o "1-1"' not in body_xml:
        f.error("J/6d", 'no table of contents field (TOC \\o "1-1")')
    for kind, block in ordered:
        if kind == "t":
            return
        if is_level1(block):
            return
        text = block.text.strip()
        if not text or "TOC" in block._p.xml:
            continue
        f.error("J/6e", "content before the first numbered section; only the "
                        "contents field belongs there — the page header "
                        "carries the title", text[:50])
        return


def check_cross_references(ordered, f):
    """F — an identifier in prose is always followed by a name."""
    cited, in_tables = set(), set()
    for kind, block in ordered:
        if kind == "t":
            # A table cell naming SOP-010 -- a Responsibilities grid, a mapping
            # table -- justifies its References row, so it counts one way only.
            # It does not oblige a new row: a Revision History entry recording
            # that GXP-OQ-001 was withdrawn must not re-list a retired document.
            # F/7 does not apply here either: a cell is a lookup key, not prose,
            # and the name is the next column over.
            for row in rows_of(block):
                for value in row:
                    in_tables.update(m.group(0) for m in ID_RE.finditer(value or ""))
            continue
        if kind != "p":
            continue
        text = block.text
        for match in ID_RE.finditer(text):
            identifier = match.group(0)
            tail = identifier.split("-", 1)[1] if "-" in identifier else ""
            if PLACEHOLDER_ID_RE.match(tail) or PLACEHOLDER_ID_RE.match(identifier):
                continue
            cited.add(identifier)
            if not RESOLVED_RE.match(text[match.end():]):
                f.error("F/7", f"{identifier} cited with no name — write "
                               f"'{identifier} (Short Name)'", text.strip()[:60])
    return cited, in_tables


def check_references(ordered, cited, in_tables, f):
    """G — a References table, and citations resolve both ways."""
    listed, in_refs, seen_table = set(), False, False
    for kind, block in ordered:
        if kind == "p":
            if is_level1(block):
                name = heading_text(block)
                if in_refs and not seen_table:
                    f.error("G/8", "References section contains no table")
                in_refs = name == "REFERENCES"
                seen_table = False
            continue
        if in_refs:
            seen_table = True
            for row in rows_of(block)[1:]:
                if row and row[0]:
                    listed.add(row[0].strip())
    if in_refs and not seen_table:
        f.error("G/8", "References section contains no table")

    if not listed:
        return
    uncited = {r for r in listed if ID_RE.fullmatch(r)} - cited - in_tables
    unlisted = cited - listed
    if unlisted:
        f.error("G/9", f"cited in the body but absent from References: "
                       f"{', '.join(sorted(unlisted))}")
    if uncited:
        f.warn("G/9", f"listed in References but never cited: "
                      f"{', '.join(sorted(uncited))}")


def check_approval_and_history(ordered, f):
    """H, I — exact columns, and Revision History newest-first."""
    current, approval, history = None, None, None
    for kind, block in ordered:
        if kind == "p":
            if is_level1(block):
                current = heading_text(block)
            continue
        if current == "APPROVAL" and approval is None:
            approval = block
        elif current == "REVISION HISTORY" and history is None:
            history = block

    if approval is None:
        f.error("H/10", "no Approval table")
    else:
        head = [c.strip() for c in rows_of(approval)[0]]
        if head != ["Role", "Name", "Title", "Signature", "Date"]:
            f.error("H/10", f"Approval columns are {head}, expected "
                            "['Role', 'Name', 'Title', 'Signature', 'Date']")

    if history is None:
        f.error("I/11", "no Revision History table")
        return
    grid = rows_of(history)
    head = [c.strip() for c in grid[0]]
    expected = ["Version", "Date", "Author", "Description of Change",
                "Change Control Ref"]
    if head != expected:
        f.error("I/11", f"Revision History columns are {head}, expected {expected}")
    versions = []
    for row in grid[1:]:
        parts = re.findall(r"\d+", row[0] if row else "")
        if parts:
            versions.append(tuple(int(p) for p in parts))
    if versions != sorted(versions, reverse=True):
        f.error("I/11", "Revision History is not newest-first")


def check_language(ordered, is_draft, f):
    """K, M — vague terms, date format, spelling, unresolved placeholders."""
    for kind, block in ordered:
        if kind != "p":
            continue
        text = block.text.strip()
        if not text:
            name = style_name(block)
            if name.startswith(("Heading", "GxP L")):
                f.error("M/12", f"empty {name} paragraph")
            continue
        lowered = text.lower()
        for term in VAGUE_TERMS:
            if re.search(rf"\b{re.escape(term)}\b", lowered):
                f.warn("K/13", f"vague term {term!r} — state something "
                               "measurable", text[:60])
        for word in UK_SPELLINGS:
            if re.search(rf"\b{word}\b", lowered):
                f.warn("K/13", f"UK spelling {word!r} (documents are US English)",
                       text[:60])
        if NON_ISO_DATE_RE.search(text):
            f.warn("K/13", "date is not ISO 8601 YYYY-MM-DD", text[:60])
        if not is_draft and PLACEHOLDER_RE.search(text):
            f.error("M/14", "unresolved placeholder in a non-draft file",
                    text[:60])


# ---------------------------------------------------------------------------


def check(path):
    path = Path(path)
    f = Findings(path)
    doc = Document(str(path))
    ordered = list(blocks(doc.element.body, doc))

    check_no_control_block(find_control_block(ordered), f)
    doc_id, version = header_identity(doc)
    if not doc_id:
        match = FILENAME_RE.match(path.name)
        doc_id = match.group("id") if match else ""
    is_draft = check_version_and_filename(path, version, doc_id, f)

    prefix = (doc_id or path.name).split("-")[0].upper()
    doc_type = PREFIX_TO_TYPE.get(prefix, "sop")

    check_header(doc, f, doc_id, version, is_draft)
    check_footer(doc, f)
    check_page_setup(doc, f)
    check_toc_and_title(ordered, doc, f)
    check_typography(doc, f)
    check_clause_ladder(doc, ordered, f)
    check_tables(ordered, f)
    check_table_candidates(ordered, f)

    sections = [heading_text(b) for k, b in ordered if k == "p" and is_level1(b)]
    check_skeleton(sections, doc_type, f)

    cited, in_tables = check_cross_references(ordered, f)
    check_references(ordered, cited, in_tables, f)
    check_approval_and_history(ordered, f)
    check_language(ordered, is_draft, f)
    return f


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("paths", nargs="+", type=Path)
    parser.add_argument("--json", action="store_true")
    parser.add_argument("--errors-only", action="store_true")
    args = parser.parse_args()

    results, failed = [], False
    for path in args.paths:
        if path.name.startswith("~$"):
            continue
        try:
            found = check(path)
        except Exception as exc:  # unreadable file is itself a finding
            results.append({"file": str(path), "findings": [
                {"level": ERROR, "rule": "-", "message": f"cannot read: {exc}",
                 "where": ""}]})
            failed = True
            continue
        items = found.errors if args.errors_only else found.items
        results.append({"file": str(path), "findings": items})
        failed = failed or bool(found.errors)

    if args.json:
        print(json.dumps(results, indent=2))
        return 1 if failed else 0

    total = 0
    for result in results:
        items = result["findings"]
        total += len(items)
        status = "clean" if not items else f"{len(items)} finding(s)"
        print(f"\n{result['file']} — {status}")
        for item in items:
            tail = f"  [{item['where']}]" if item["where"] else ""
            print(f"  {item['level']:<7} {item['rule']:<7} {item['message']}{tail}")
    print(f"\n{total} finding(s) across {len(results)} file(s)")
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
