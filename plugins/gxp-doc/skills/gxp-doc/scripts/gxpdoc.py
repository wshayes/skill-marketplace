"""Builder for GxP controlled documents.

The single source of document styling. Every rule implemented here is
specified in ../references/style-guide.md; if the two disagree, the style
guide is wrong and should be corrected to match observable output.

Run under:  uv run --with python-docx python <script>
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

FONT = "Times New Roman"
TABLE_HEADER_FILL = "F2F2F2"
TABLE_BORDER = "D9D9D9"
HEADER_BORDER = "000000"

#: Resolved per project. SKILL.md tells the agent to read the real name from
#: memory and ask when it is not there; the placeholder must never ship.
COMPANY_PLACEHOLDER = "<Company Name>"
COMPANY = COMPANY_PLACEHOLDER

#: doc_type -> the all-caps banner across the top of the page header
BANNER = {
    "qm": "QUALITY MANUAL",
    "policy": "QUALITY POLICY",
    "sop": "STANDARD OPERATING PROCEDURE",
    "wi": "WORK INSTRUCTION",
    "form": "FORM",
    "log": "LOG",
    "protocol": "VALIDATION PROTOCOL",
    "report": "VALIDATION REPORT",
}

#: clause level -> (number x, text x) in inches. Hanging indent at every
#: level so wrapped lines align to the text column, never the number.
INDENT = {
    1: (0.00, 0.30),
    2: (0.30, 0.65),
    3: (0.65, 1.10),
    4: (1.10, 1.65),
}
BULLET_INDENT = (1.65, 1.90)
BODY_INDENT = 0.30

#: Rows that must never appear in a control block. A document carries no
#: date and no status; those are register columns on the Document Master List.
FORBIDDEN_CONTROL_ROWS = {
    "status",
    "date",
    "effective date",
    "approval date",
    "last reviewed",
    "review date",
}

#: Rows seen in the legacy control blocks this style guide removes. Kept only
#: so check_doc and restyle_doc can recognise such a table and strip it.
LEGACY_CONTROL_ROWS = [
    "Company", "Document ID", "Title", "Document Type", "Version",
    "Review Cycle", "Owner", "Approver", "Regulatory Basis", "Supersedes",
]

FILENAME_RE = re.compile(
    r"^(?P<id>[A-Z]{2,4}(?:-[A-Z]{2,5}){0,2}-\d{2,4}(?:-\d{3})?)"
    r"_(?P<slug>[A-Za-z0-9_]+)"
    r"_v(?P<version>\d+\.\d+)"
    r"(?:_DRAFT\d+)?\.docx$"
)
VERSION_RE = re.compile(r"^\d+\.\d+$")


# --------------------------------------------------------------------------
# OXML helpers -- the parts python-docx will not do for us
# --------------------------------------------------------------------------


def _inches(value):
    """Inches snapped to whole twips.

    Word stores indents in twips; a raw float lands a stray EMU off and a
    value written then read back no longer compares equal to its source.
    """
    return Emu(int(round(value * 1440)) * 635)


def _version_key(row):
    """Sort key tolerant of legacy version cells like "2 (pending)"."""
    parts = re.findall(r"\d+", str(row[0]))
    return tuple(int(p) for p in parts) or (0,)


def _set_font(run, size=11, bold=False, italic=False, caps=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.all_caps = caps
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)
    return run


def _field(paragraph, instruction, placeholder="1", size=9):
    """Insert a live Word field (PAGE, NUMPAGES, TOC ...).

    python-docx has no field support; without this, "Page X of Y" is frozen
    text -- which is exactly why every current QMS document prints "Page of".
    """
    begin = paragraph.add_run()
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "begin")
    begin._r.append(el)

    instr = paragraph.add_run()
    el = OxmlElement("w:instrText")
    el.set(qn("xml:space"), "preserve")
    el.text = f" {instruction} "
    instr._r.append(el)

    sep = paragraph.add_run()
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "separate")
    sep._r.append(el)

    result = paragraph.add_run(placeholder)

    end = paragraph.add_run()
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "end")
    end._r.append(el)

    for run in (begin, instr, sep, result, end):
        _set_font(run, size=size)
    return result


def _outline_level(style, level):
    """Give a clause style an outline level, zero-based.

    Word's contents field collects by outline level, not by style name: a
    ``TOC \\o "1-1"`` field gathers every paragraph at outline level 1. A
    custom style based on Normal inherits outline level "body text", so
    without this the GxP L1 headings are invisible to the field and the
    contents come back empty -- and the Navigation Pane stays flat. Setting it
    on the style rather than the paragraph keeps the numbers literal text
    (J/6c) while still making the structure machine-readable.
    """
    ppr = style.element.get_or_add_pPr()
    for existing in ppr.findall(qn("w:outlineLvl")):
        ppr.remove(existing)
    node = OxmlElement("w:outlineLvl")
    node.set(qn("w:val"), str(level - 1))
    ppr.append(node)


def _shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _borders(table, outer_eighths, inner_eighths, color):
    """Border widths are in eighths of a point: 12 == 1.5pt, 6 == 0.75pt."""
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    spec = (
        ("top", outer_eighths),
        ("left", outer_eighths),
        ("bottom", outer_eighths),
        ("right", outer_eighths),
        ("insideH", inner_eighths),
        ("insideV", inner_eighths),
    )
    for tag, size in spec:
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def _repeat_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def _no_split(row):
    """Keep a row's content on one page.

    Without this a long cell breaks mid-row and the neighbouring cells are
    left blank on the continuation page, which reads as missing data.
    """
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def _cell_margins(table, inches=0.08):
    tbl_pr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for tag in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(int(inches * 1440)))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.append(mar)


def _merge_row(table, row_index):
    row = table.rows[row_index]
    merged = row.cells[0]
    for cell in row.cells[1:]:
        merged = merged.merge(cell)
    return merged


def _write_cell(cell, text, size=10, bold=False, caps=False, align=None):
    cell.text = ""
    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(text)
    _set_font(run, size=size, bold=bold, caps=caps)
    return para


# --------------------------------------------------------------------------
# Builder
# --------------------------------------------------------------------------


class GxpDoc:
    """Build one controlled document.

    Note the constructor takes no ``status`` and no ``date``: a document
    carries neither. It is a draft until it is signed, and its approved date
    is the last signature date in the Approval table.
    """

    def __init__(self, doc_id, title, doc_type, version, company=COMPANY,
                 logo_path=None):
        if doc_type not in BANNER:
            raise ValueError(
                f"unknown doc_type {doc_type!r}; expected one of {sorted(BANNER)}"
            )
        if not VERSION_RE.match(str(version)):
            raise ValueError(
                f"version must be X.Y, got {version!r}. The version field never "
                "reads DRAFT -- draft state belongs in the filename."
            )
        self.doc_id = doc_id
        self.title = title
        self.doc_type = doc_type
        self.version = str(version)
        self.company = company
        self.logo_path = logo_path

        self.doc = Document()
        self._counters = [0, 0, 0, 0]
        self._section_names = []

        self._page_setup()
        self._update_fields_on_open()
        self._define_styles()
        self._build_header()
        self._build_footer()

    # -- setup ------------------------------------------------------------

    def _page_setup(self):
        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        for side in ("left", "right", "top", "bottom"):
            setattr(section, f"{side}_margin", Inches(1))
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    def _update_fields_on_open(self):
        """Tell Word to compute every field when the document is opened.

        The contents field and the footer's PAGE/NUMPAGES are real fields, not
        frozen text -- that is the whole point of B/2 and J/6d. But a field
        carries a cached result, and python-docx cannot compute one: it has no
        layout engine, so it cannot know what page a heading lands on. Without
        this flag the cache is the placeholder, and every document opens
        reading "Contents -- update this field to populate" and "Page 1 of 1"
        until somebody presses F9. Setting w:updateFields makes Word refresh
        them on open, which is the only place the real values can come from.
        """
        settings = self.doc.settings.element
        for existing in settings.findall(qn("w:updateFields")):
            settings.remove(existing)
        flag = OxmlElement("w:updateFields")
        flag.set(qn("w:val"), "true")
        settings.append(flag)

    def _define_styles(self):
        styles = self.doc.styles

        normal = styles["Normal"]
        normal.font.name = FONT
        normal.font.size = Pt(11)
        normal.font.color.rgb = RGBColor(0, 0, 0)
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), FONT)
        normal.paragraph_format.line_spacing = 1.15
        normal.paragraph_format.space_after = Pt(6)

        body = styles.add_style("GxP Body", WD_STYLE_TYPE.PARAGRAPH)
        body.base_style = styles["Normal"]
        body.paragraph_format.left_indent = _inches(BODY_INDENT)

        for level, (number_x, text_x) in INDENT.items():
            style = styles.add_style(f"GxP L{level}", WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles["Normal"]
            fmt = style.paragraph_format
            fmt.left_indent = _inches(text_x)
            fmt.first_line_indent = _inches(number_x - text_x)
            fmt.tab_stops.add_tab_stop(_inches(text_x), WD_TAB_ALIGNMENT.LEFT)
            _outline_level(style, level)
            if level == 1:
                style.font.size = Pt(13)
                style.font.bold = True
                style.font.all_caps = True
                fmt.space_before = Pt(12)
                fmt.keep_with_next = True
            else:
                style.font.size = Pt(11)
                style.font.bold = False

        bullet = styles.add_style("GxP Bullet", WD_STYLE_TYPE.PARAGRAPH)
        bullet.base_style = styles["Normal"]
        fmt = bullet.paragraph_format
        fmt.left_indent = _inches(BULLET_INDENT[1])
        fmt.first_line_indent = _inches(BULLET_INDENT[0] - BULLET_INDENT[1])
        fmt.tab_stops.add_tab_stop(_inches(BULLET_INDENT[1]), WD_TAB_ALIGNMENT.LEFT)

    # -- page furniture ---------------------------------------------------

    def _build_header(self):
        header = self.doc.sections[0].header
        wordmark = header.paragraphs[0]
        wordmark.alignment = WD_ALIGN_PARAGRAPH.LEFT
        wordmark.paragraph_format.space_after = Pt(2)
        if self.logo_path:
            wordmark.add_run().add_picture(str(self.logo_path), height=Inches(0.35))
        else:
            _set_font(wordmark.add_run(self.company), size=14, bold=True)

        table = header.add_table(rows=3, cols=3, width=Inches(7.0))
        table.autofit = False
        for row in table.rows:
            for cell, width in zip(row.cells, (1.5, 4.2, 1.3)):
                cell.width = Inches(width)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        centre = WD_ALIGN_PARAGRAPH.CENTER
        banner = _merge_row(table, 0)
        _write_cell(banner, BANNER[self.doc_type], size=12, bold=True,
                    caps=True, align=centre)

        for cell, label in zip(
            table.rows[1].cells,
            ("Document ID", "Document Description", "Version No."),
        ):
            _write_cell(cell, label, size=10, align=centre)

        values = (self.doc_id, self.title.upper(), self.version)
        for cell, value in zip(table.rows[2].cells, values):
            _write_cell(cell, value, size=10, bold=True, align=centre)

        # The header block is the one table with solid black rules and no
        # shading -- it reads as furniture, not content.
        _borders(table, outer_eighths=12, inner_eighths=6, color=HEADER_BORDER)
        _cell_margins(table, 0.05)
        header.add_paragraph()

    def _build_footer(self):
        footer = self.doc.sections[0].footer
        para = footer.paragraphs[0]
        para.paragraph_format.tab_stops.add_tab_stop(
            Inches(3.25), WD_TAB_ALIGNMENT.CENTER
        )
        para.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.5), WD_TAB_ALIGNMENT.RIGHT
        )
        _set_font(para.add_run("Printed copies are uncontrolled\t"), size=9)
        _set_font(para.add_run("Page "), size=9)
        _field(para, "PAGE", "1")
        _set_font(para.add_run(" of "), size=9)
        _field(para, "NUMPAGES", "1")

    # -- body -------------------------------------------------------------

    def toc(self):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _field(para, 'TOC \\o "1-1" \\h \\z \\u',
               "Contents — update this field to populate.", size=11)
        self.doc.add_paragraph()
        return para

    def h1(self, text):
        self._counters[0] += 1
        self._counters[1:] = [0, 0, 0]
        number = f"{self._counters[0]}."
        self._section_names.append(text.upper())
        return self._numbered(1, number, text.upper())

    def clause(self, level, text):
        if not 2 <= level <= 4:
            raise ValueError("clause level must be 2, 3 or 4; use h1() for 1")
        if self._counters[0] == 0:
            raise ValueError("a clause must follow an h1()")
        self._counters[level - 1] += 1
        for deeper in range(level, 4):
            self._counters[deeper] = 0
        number = ".".join(str(c) for c in self._counters[:level]) + "."
        return self._numbered(level, number, text)

    def _numbered(self, level, number, text):
        para = self.doc.add_paragraph(style=f"GxP L{level}")
        size, bold = (13, True) if level == 1 else (11, False)
        _set_font(para.add_run(f"{number}\t"), size=size, bold=bold)
        _set_font(para.add_run(text), size=size, bold=bold)
        return para

    def body(self, text):
        para = self.doc.add_paragraph(style="GxP Body")
        _set_font(para.add_run(text), size=11)
        return para

    def bullet(self, text):
        para = self.doc.add_paragraph(style="GxP Bullet")
        _set_font(para.add_run("•\t"), size=11)
        _set_font(para.add_run(text), size=11)
        return para

    # -- tables -----------------------------------------------------------

    def table(self, headers, rows, widths=None, fillable=False):
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.autofit = widths is None

        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            _write_cell(cell, str(header), size=10, bold=True)
            _shade(cell, TABLE_HEADER_FILL)
        _repeat_header_row(table.rows[0])

        for row_index, row in enumerate(rows, start=1):
            for col_index, value in enumerate(row):
                _write_cell(table.rows[row_index].cells[col_index],
                            "" if value is None else str(value), size=10)
            if fillable:
                table.rows[row_index].height = Inches(0.30)

        for row in table.rows:
            _no_split(row)

        if widths:
            for row in table.rows:
                for cell, width in zip(row.cells, widths):
                    cell.width = Inches(width)

        _borders(table, outer_eighths=6, inner_eighths=4, color=TABLE_BORDER)
        _cell_margins(table)
        self.doc.add_paragraph()
        return table

    def references(self, groups):
        """groups: {group name: [(document id, title), ...]}.

        Always a table, never a comma-delimited paragraph. Grouped into
        numbered clauses once there is more than one group.
        """
        self.h1("References")
        multiple = len(groups) > 1
        for name, entries in groups.items():
            if multiple:
                self.clause(2, name)
            self.table(["Document ID", "Title"], entries, widths=(1.8, 4.9))
        return self._section_names[-1]

    def revision_history(self, rows):
        """rows: [(version, date, author, description, change control ref)].

        Rendered newest-first -- the current revision is the interesting one.
        """
        self.h1("Revision History")
        ordered = sorted(rows, key=_version_key, reverse=True)
        return self.table(
            ["Version", "Date", "Author", "Description of Change",
             "Change Control Ref"],
            ordered,
            widths=(0.8, 0.9, 1.1, 3.0, 1.2),
        )

    def approval(self, signatories):
        """signatories: [(role, name, title)]. Date stays blank until signed."""
        self.h1("Approval")
        self.body(
            "This document is approved by the signatories below. Wet-ink "
            "signatures may be replaced by GxPSign qualified electronic "
            "signatures applied through the platform; the electronic signature "
            "block, capturing signer identity, meaning, and timestamp, occupies "
            "the signature space provided."
        )
        rows = [(role, name, title, "", "") for role, name, title in signatories]
        table = self.table(
            ["Role", "Name", "Title", "Signature", "Date"],
            rows,
            widths=(1.0, 1.6, 1.6, 1.7, 1.0),
        )
        for row in table.rows[1:]:
            row.height = Inches(0.40)
        return table

    # -- output -----------------------------------------------------------

    def save(self, path):
        path = Path(path)
        match = FILENAME_RE.match(path.name)
        if not match:
            raise ValueError(
                f"{path.name!r} does not match "
                "<ID>_<Title_Snake>_v<X.Y>[_DRAFT<N>].docx"
            )
        if match.group("version") != self.version:
            raise ValueError(
                f"filename says v{match.group('version')} but the document "
                f"says {self.version}"
            )
        if match.group("id") != self.doc_id:
            raise ValueError(
                f"filename says {match.group('id')} but the document says "
                f"{self.doc_id}"
            )
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))
        return path
